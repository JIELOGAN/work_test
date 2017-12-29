#!/usr/bin/env python
# _#_ coding:utf-8 _*_

from django.contrib.auth.decorators import login_required
from django.shortcuts import render_to_response, render, redirect
from django.http import HttpResponse, HttpResponseNotFound, JsonResponse
from django import forms
from work_app.models import *
from django.contrib import auth
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.decorators import api_view
from rest_framework import status
from django.http import Http404
from django.http import HttpResponseRedirect
# from pure_pagination import Paginator, EmptyPage, PageNotAnInteger
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_DIRECTION


@login_required(login_url='/login')
@api_view(['GET', 'POST'])
def file_parse(request):
    if request.method =='GET':
        assets_info = sinorail_assets.objects.all()
        station_category = sinorail_station.objects.all()
        station_count = station_category.count()
## 筛选车站
        station_id = request.GET.get('station', "")
        if station_id:
            assets_info = sinorail_assets.objects.filter(station_id=int(station_id))
        assets_count = assets_info.count()
## 分页
        page = request.GET.get('page', 1)
        paginator = Paginator(assets_info, 8)
        try:
            assets_page = paginator.page(page)
        except PageNotAnInteger:
            assets_page = paginator.page(1)
        except EmptyPage:
            assets_page = paginator.page(paginator.num_pages)

        context = {'assets_info': assets_info, 'assets_page': assets_page, 'station_category': station_category,
                   'station_count': station_count, 'station_id': station_id, 'assets_count': assets_count}
        return render(request, 'logfile/file_parse.html', context)

    elif request.method == 'POST':
        button_name = request.data.get('button', None)  #获取数据，rest框架中数据存在request.data中
        button_data = request.data.get('data', None)
        if button_name == 'button1':
            print(button_name+'----'+button_data)
            # return render(request, 'logfile/file_parse.html', {})
            return HttpResponse({'status': 'success'}, content_type='application/json')
        # else:
        #     return HttpResponse("{'status':'fail', 'msg':'失败'}", content_type='application/json')

        elif button_name == 'button2':
            try:
                # hardware_data = hardware_info.objects.filter(sn=button_data).values('sn', 'slot_status', 'power_state',
                #                                                                     'fan_state', 'device_temperature',
                #                                                                     'card_UsingPercent', 'vender_name',
                #                                                                     'cpu_usage', 'memory_usage')
                station_id = sinorail_assets.objects.filter(sn=button_data).values('station_id')
                print(station_id[0].values())
                station_name = sinorail_station.objects.filter(id=station_id).values('station_name')
                station_name = station_name[0].values()
                print(station_name)
                system_name = source_system.objects.filter(station_id_id=station_id).values('system_name')
                line_id = sinorail_station.objects.filter(station_name=station_name).values('line_id_id')
                line_name = train_line.objects.filter(id=line_id).values('line_name')

                assets_data = sinorail_assets.objects.filter(station_id=station_id).values('id','sn','device_name','assets_type','ip_address','cab_address')
                assets_length = len(sinorail_assets.objects.filter(station_id=station_id))

                report = Document()  # 创建空白文档
                ## 新建样式
                style_T1 = report.styles.add_style('T1', WD_STYLE_TYPE.PARAGRAPH)
                style_T1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                style_T1.paragraph_format.page_break_before = True
                style_T1.paragraph_format.widow_control = True
                style_H1 = report.styles.add_style('H1', WD_STYLE_TYPE.PARAGRAPH)
                style_H1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                style_H1.paragraph_format.widow_control = True
                style_H2 = report.styles.add_style('H2', WD_STYLE_TYPE.PARAGRAPH)
                style_H2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                style_H2.paragraph_format.widow_control = True
                # head1 = report.add_heading('巡检报告', 0)
                head1 = report.add_paragraph(u'巡检报告', 'T1')
                # run = head1.add_run(u'巡检报告')
                # run.font.name = u'黑体'
                # r = run._element
                # r.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                head1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
                p_day = report.add_paragraph('日期：     年     月   日', 'Normal')
                p_day.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                paragraph1 = report.add_paragraph('站点：', 'Normal')
                paragraph1.add_run(station_name)
                # paragraph2 = report.add_paragraph('线路：')
                # paragraph2.add_run(str(line_name))
                # paragraph3 = report.add_paragraph('系统：')
                # paragraph3.add_run(str(system_name))

                report.add_paragraph('1.硬件清单', 'H1')
                report.add_paragraph('网络设备：共 '+ str(assets_length) + ' 台', 'H2')

                # 创建表
                device_table = report.add_table(assets_length + 1, 6, style='Table Grid')
                report.add_paragraph('本期维修项目：', 'Normal')
                report.add_paragraph('_____________________________________________________________')
                report.add_paragraph('本期检查项目：', 'Normal')
                report.add_paragraph('_____________________________________________________________')
                report.add_paragraph('一、设备状态检查', 'H2')
                report.add_paragraph('1.硬件信息检查', 'Normal')
                hardware_table = report.add_table(5, 2, style='Table Grid')
                hardware_table.cell(0, 0).text = 'CPU                [ ]正常           [ ]不正常'
                hardware_table.cell(0, 1).text = '内存               [ ]正常           [ ]不正常'
                hardware_table.cell(1, 0).text = '引擎同步状态       [ ]正常           [ ]不正常'
                hardware_table.cell(1, 1).text = '系统日志           [ ]正常           [ ]不正常'
                hardware_table.cell(2, 0).text = 'Flash利用率        [ ]正常           [ ]不正常'
                hardware_table.cell(2, 1).text = '电源是否正常       [ ]正常           [ ]不正常'
                hardware_table.cell(3, 0).text = '风扇模块           [ ]正常           [ ]不正常'
                hardware_table.cell(3, 1).text = '单板运作状态       [ ]正常           [ ]不正常'
                hardware_table.cell(4, 0).text = '光模块信息         [ ]正常           [ ]不正常'
                hardware_table.cell(4, 1).text = '设备温度           [ ]正常           [ ]不正常'
                report.add_paragraph('异常情况说明：_____________________________________________________________', 'Normal')
                report.add_paragraph('2.运行状态检查', 'Normal')
                runtime_table = report.add_table(2, 2, style='Table Grid')
                runtime_table.cell(0, 0).text = '系统时钟（NTP）     [ ]正常           [ ]不正常'
                runtime_table.cell(0, 1).text = '启动文件信息         [ ]正常          [ ]不正常'
                runtime_table.cell(1, 0).text = 'OSPF错误统计        [ ]正常           [ ]不正常'
                runtime_table.cell(1, 1).text = 'OSPF邻居信息         [ ]正常          [ ]不正常'
                report.add_paragraph('异常情况说明：_____________________________________________________________', 'Normal')
                report.add_paragraph('3.安全性检查', 'Normal')
                security_table = report.add_table(2, 2, style='Table Grid')
                security_table.cell(0, 0).text = '登录用户口令安全性       [ ]正常      [ ]不正常'
                security_table.cell(0, 1).text = 'VTY用户界面安全性       [ ]正常      [ ]不正常'
                security_table.cell(1, 0).text = 'FTP服务状态             [ ]正常      [ ]不正常'
                security_table.cell(1, 1).text = 'SNMP安全性              [ ]正常      [ ]不正常'
                report.add_paragraph('异常情况说明：_____________________________________________________________', 'Normal')
                report.add_page_break()
                report.add_paragraph('本期巡检工作总结：', 'H2')
                report.add_paragraph('本次预防性发现的问题：_____________________________________________________________', 'Normal')
                report.add_paragraph('本次巡检解决的问题：_____________________________________________________________', 'Normal')
                report.add_paragraph('其他备注：_____________________________________________________________', 'Normal')
                report.add_paragraph('客户签名：', 'Normal')
                report.add_paragraph('信息技术所                                             联创中心', 'Normal')
                report.add_paragraph('工程师签名：__________                                 工程师签名：__________', 'Normal')
                report.add_paragraph('日     期：__________                                  日     期：__________', 'Normal')

                # report.add_page_break()
                # report.add_heading('参数说明:', 6)
                # unusual_table = report.add_table(len(hardware_data[0]) + 1, 3, style='Table Grid')

                # 获取表头
                ## 网络设备表
                heading_cells = device_table.rows[0].cells
                keys = assets_data[0].keys()
                keys = list(keys)
                for r in range(len(keys)):
                    heading_cells[r].text = keys[r]
                ## 异常说明表
                # unusual_table_heading_cells = unusual_table.rows[0].cells
                # unusual_table_heading_cells[0].text = '参数名'
                # unusual_table_heading_cells[1].text = '参数说明'
                # unusual_table_heading_cells[2].text = '是否正常'

                # 获取表值
                for i in range(assets_length):
                    values = list(assets_data[i].values())
                    for r in range(len(keys)):
                        device_table.cell(i + 1, r).text = str(values[r])

                # 添加一列序号
                device_table.add_column(Inches(0.25))
                device_table.cell(0, 6).text = 'num'
                for i in range(assets_length):
                    device_table.cell(i+1,6).text = str(i+1)

                # 填入状态说明
                # temp_data = list(hardware_data[0].items())
                # for i in range(len(temp_data)):
                #     unusual_table.cell(i + 1, 0).text = str(temp_data[i][0])
                #     unusual_table.cell(i + 1, 1).text = str(temp_data[i][1])

                # 列宽
                # for r in range(6):
                #     device_table.autofit = False
                #     width = float(len(heading_cells[r].text)) / 5
                #     print(width)
                #     print(Inches(width))
                #     device_table.columns[r].width = Inches(width)
                # 字体
                report.styles['T1'].font.name = u'黑体'
                report.styles['T1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                report.styles['T1'].font.size = Pt(18)
                report.styles['H1'].font.name = u'宋体'
                report.styles['H1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                report.styles['H1'].font.size = Pt(12)
                report.styles['H2'].font.name = u'宋体'
                report.styles['H2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                report.styles['H2'].font.size = Pt(10.5)
                report.styles['Normal'].font.name = u'宋体'
                report.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                report.styles['Normal'].font.size = Pt(9)
                def table_font(table):
                    for row in table.rows:
                        for cell in row.cells:
                            pargraphs = cell.paragraphs
                            for pargraph in pargraphs:
                                for run in pargraph.runs:
                                    font = run.font
                                    font.size = Pt(7.5)
                # def table_font(table):
                #     table.style.font.size = Pt(7.5)
                table_font(device_table)
                table_font(hardware_table)
                table_font(runtime_table)
                table_font(security_table)
                device_table.autofit = True
                device_table.direction = WD_TABLE_DIRECTION.LTR
                report.save(button_data + 'sinorail_report.docx')
                # return HttpResponse({'status': 'success'}, content_type='application/json')
                return JsonResponse({'status': 'success'})

            except IndexError:
                print('没有该设备的日志文件')
                return JsonResponse({'status': 'fail'})
            # return JsonResponse({'status': 'success'})
    # else:
    #     return render(request, 'logfile/file_parse.html', {})
